"""
DOCX 문서 생성 모듈 v2
- 표 안 인라인 서식 지원
- 페이지 분리 개선
- 표 너비/정렬 개선
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

from .styles import DocumentStyle, DEFAULT_STYLE
from .parser import (
    DocumentMetadata, ParsedBlock, BlockType,
    ParsedTable, ParsedFlow, MarkdownParser
)


class DocxGenerator:
    """DOCX 문서 생성기"""

    def __init__(self, style: DocumentStyle = None):
        self.style = style or DEFAULT_STYLE
        self.doc = Document()
        self.heading_counters = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
        self.is_first_h1 = True  # 첫 H1인지 추적
        self._content_width = self.style.page.content_width  # 동적 콘텐츠 너비
        self._setup_document()

    def _setup_document(self):
        """문서 초기 설정"""
        for section in self.doc.sections:
            section.page_width = self.style.page.width
            section.page_height = self.style.page.height
            section.top_margin = self.style.page.margin_top
            section.bottom_margin = self.style.page.margin_bottom
            section.left_margin = self.style.page.margin_left
            section.right_margin = self.style.page.margin_right

        self._setup_styles()

    def _setup_styles(self):
        """문서 스타일 설정"""
        styles = self.doc.styles

        normal = styles['Normal']
        normal.font.name = self.style.fonts.default_font
        normal.font.size = self.style.fonts.body_size
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), self.style.fonts.default_font)

        # 제목은 _render_heading에서 일반 단락 + 직접 서식으로 처리
        # (Word 내장 Heading 스타일의 불릿/기호 문제 방지)

    def generate(self, metadata: Optional[DocumentMetadata], blocks: List[ParsedBlock]) -> Document:
        """문서 생성"""
        if metadata:
            self._create_cover_page(metadata)
            self._add_page_break()

        # 목차 페이지 제거 (필요시 수동 추가)
        # self._create_toc()
        # self._add_page_break()

        for block in blocks:
            self._render_block(block)

        return self.doc

    def save(self, path: str):
        """문서 저장"""
        self.doc.save(path)

    # =========================================================================
    # 표지 생성
    # =========================================================================

    def _create_cover_page(self, meta: DocumentMetadata):
        """표지 페이지 생성"""
        for _ in range(4):
            self.doc.add_paragraph()

        if meta.logo and Path(meta.logo).exists():
            logo_para = self.doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(meta.logo, width=Inches(2))
            self.doc.add_paragraph()

        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(meta.title)
        title_run.bold = True
        title_run.font.size = self.style.fonts.title_size
        title_run.font.name = self.style.fonts.default_font

        if meta.subtitle:
            self.doc.add_paragraph()
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(meta.subtitle)
            subtitle_run.bold = True
            subtitle_run.font.size = self.style.fonts.subtitle_size
            subtitle_run.font.name = self.style.fonts.default_font

        if meta.description:
            self.doc.add_paragraph()
            desc_para = self.doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_para.add_run(meta.description)
            desc_run.font.size = self.style.fonts.body_size
            desc_run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.text_gray)

        for _ in range(4):
            self.doc.add_paragraph()

        self._create_meta_table(meta)

        for _ in range(4):
            self.doc.add_paragraph()

        if meta.confidential:
            conf_para = self.doc.add_paragraph()
            conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            conf_run = conf_para.add_run(self.style.confidential_text)
            conf_run.font.size = self.style.fonts.small_size
            conf_run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.text_light_gray)

    def _create_meta_table(self, meta: DocumentMetadata):
        """메타 정보 테이블 생성"""
        meta_items = [
            ('프로젝트', meta.project),
            ('작성일', meta.date),
            ('버전', meta.version),
            ('작성', meta.author),
        ]
        meta_items = [(k, v) for k, v in meta_items if v]

        if not meta_items:
            return

        table = self.doc.add_table(rows=len(meta_items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(3.0)

        for i, (label, value) in enumerate(meta_items):
            label_cell = table.cell(i, 0)
            label_cell.text = label
            self._style_cell(label_cell, bg_color=self.style.colors.bg_info, bold=True, align='right')
            label_cell.width = Inches(1.5)

            value_cell = table.cell(i, 1)
            value_cell.text = value
            self._style_cell(value_cell, align='left')
            value_cell.width = Inches(3.0)

        self._set_table_borders(table)

    # =========================================================================
    # 목차 생성
    # =========================================================================

    def _create_toc(self):
        """목차 생성"""
        toc_title = self.doc.add_paragraph(self.style.toc_title)
        toc_title.style = 'Heading 1'

        para = self.doc.add_paragraph()
        run = para.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        note = self.doc.add_paragraph()
        note_run = note.add_run('(목차를 업데이트하려면 Word에서 F9를 누르세요)')
        note_run.font.size = self.style.fonts.caption_size
        note_run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.text_light_gray)
        note_run.italic = True

    # =========================================================================
    # 블록 렌더링
    # =========================================================================

    def _render_block(self, block: ParsedBlock):
        """블록 렌더링"""
        if block.type == BlockType.HEADING:
            self._render_heading(block.content, block.level)
        elif block.type == BlockType.PARAGRAPH:
            self._render_paragraph(block.content)
        elif block.type == BlockType.LIST:
            self._render_list(block.content)
        elif block.type == BlockType.TABLE:
            self._render_table(block.content)
        elif block.type == BlockType.CODE:
            self._render_code(block.content)
        elif block.type == BlockType.QUOTE:
            self._render_quote(block.content)
        elif block.type == BlockType.CALLOUT_INFO:
            self._render_callout(block.content, 'info')
        elif block.type == BlockType.CALLOUT_WARNING:
            self._render_callout(block.content, 'warning')
        elif block.type == BlockType.FLOW:
            self._render_flow(block.content)
        elif block.type == BlockType.HORIZONTAL_RULE:
            self._render_hr()
        elif block.type == BlockType.PAGE_BREAK:
            self._add_page_break()
        elif block.type == BlockType.IMAGE:
            self._render_image(block.content)

    # 레벨별 폰트 크기 매핑
    HEADING_FONT_SIZES = {
        1: None,  # style.fonts.h1_size (24pt)
        2: None,  # style.fonts.h2_size (18pt)
        3: None,  # style.fonts.h3_size (14pt)
        4: Pt(12),
        5: Pt(11),
        6: Pt(10),
    }

    def _render_heading(self, text: str, level: int):
        """제목 렌더링"""
        # H1 앞에 페이지 나누기 (첫 번째 제외)
        if level == 1:
            if not self.is_first_h1:
                self._add_page_break()
            self.is_first_h1 = False

        # 번호 업데이트
        self.heading_counters[level] += 1
        for l in range(level + 1, 7):
            self.heading_counters[l] = 0

        # 번호 문자열 생성
        if self.style.heading_numbering:
            if level == 1:
                number = f"{self.heading_counters[1]}."
            elif level == 2:
                number = f"{self.heading_counters[1]}.{self.heading_counters[2]}"
            elif level == 3:
                number = f"{self.heading_counters[1]}.{self.heading_counters[2]}.{self.heading_counters[3]}"
            else:
                number = ""
            full_text = f"{number} {text}"
        else:
            full_text = text

        # 일반 단락으로 생성 후 직접 서식 적용 (내장 Heading 스타일의 기호 문제 방지)
        para = self.doc.add_paragraph()
        run = para.add_run(full_text)
        run.bold = True
        run.font.name = self.style.fonts.default_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style.fonts.default_font)
        run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.text_black)

        # 레벨별 폰트 크기 (H1~H3: 스타일 정의, H4~H6: 고정값)
        font_size = {
            1: self.style.fonts.h1_size,   # 24pt
            2: self.style.fonts.h2_size,   # 18pt
            3: self.style.fonts.h3_size,   # 14pt
        }.get(level, self.HEADING_FONT_SIZES.get(level, Pt(11)))
        run.font.size = font_size

        # 레벨별 간격
        spacing = self.style.spacing
        spacing_map = {
            1: (spacing.h1_before, spacing.h1_after),
            2: (spacing.h2_before, spacing.h2_after),
            3: (spacing.h3_before, spacing.h3_after),
            4: (Pt(10), Pt(5)),
            5: (Pt(8), Pt(4)),
            6: (Pt(6), Pt(3)),
        }
        before, after = spacing_map.get(level, (Pt(6), Pt(3)))
        para.paragraph_format.space_before = before
        para.paragraph_format.space_after = after

    def _render_paragraph(self, text: str):
        """단락 렌더링"""
        para = self.doc.add_paragraph()
        self._add_formatted_text(para, text)
        para.paragraph_format.space_after = self.style.spacing.para_after

    def _render_list(self, items: List[Dict]):
        """목록 렌더링 (일반 단락 + 수동 불릿/번호로 완전 제어)"""
        for item in items:
            para = self.doc.add_paragraph()
            level = item.get('level', 0)
            is_numbered = item['type'] != 'bullet'

            # 들여쓰기 계산
            base_indent = 0.3   # 기본 들여쓰기 (inches)
            level_step = 0.25   # 레벨당 추가 들여쓰기
            hang = 0.2          # 행잉 인덴트 (불릿/번호 위치)

            if is_numbered:
                # 번호 목록: 항상 고정 인덴트 (소스 들여쓰기 무시)
                text_indent = base_indent
            else:
                # 불릿: 번호 아이템의 하위면 1단계, 독립이면 레벨 기반
                if level >= 2:
                    # 번호 목록 하위 불릿 (indent 4칸 = level 2) → 1단계로 정규화
                    text_indent = base_indent + level_step
                else:
                    text_indent = base_indent + level_step * level

            para.paragraph_format.left_indent = Inches(text_indent)
            para.paragraph_format.first_line_indent = Inches(-hang)

            if not is_numbered:
                # 레벨별 불릿 문자 차별화
                bullet_char = '•' if level <= 1 else '-'
                prefix_run = para.add_run(f"{bullet_char} ")
            else:
                # 번호 목록
                number = item.get('number', '')
                prefix_run = para.add_run(f"{number}. " if number else "")

            prefix_run.font.name = self.style.fonts.default_font
            prefix_run.font.size = self.style.fonts.body_size
            prefix_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style.fonts.default_font)

            # 콘텐츠 (인라인 서식 적용)
            self._add_formatted_text(para, item['content'])
            para.paragraph_format.space_after = Pt(3)

    def _render_table(self, table: ParsedTable):
        """표 렌더링 (개선됨)"""
        rows = len(table.rows) + 1
        cols = len(table.headers)

        doc_table = self.doc.add_table(rows=rows, cols=cols)
        doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc_table.autofit = False

        # 페이지 콘텐츠 너비 계산 (A4 - 여백, 동적)
        content_width_emu = int(self._content_width)
        col_width = int(content_width_emu / cols)

        # 열 너비 균등 분배
        for col in doc_table.columns:
            col.width = col_width

        # 헤더 행
        for i, header in enumerate(table.headers):
            cell = doc_table.cell(0, i)
            cell.width = col_width

            # 셀 내용 - 정렬 먼저 설정 후 콘텐츠 추가
            para = cell.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_formatted_text(para, header)

            self._style_cell(
                cell,
                bg_color=self.style.table.header_bg,
                bold=True,
                align='center'
            )

        # 데이터 행 - 모든 셀 중앙 정렬 통일
        for row_idx, row in enumerate(table.rows, 1):
            for col_idx, cell_text in enumerate(row):
                if col_idx < cols:
                    cell = doc_table.cell(row_idx, col_idx)
                    cell.width = col_width

                    # 정렬 먼저 설정 후 콘텐츠 추가
                    para = cell.paragraphs[0]
                    para.clear()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._add_formatted_text(para, cell_text)

                    self._style_cell(cell, align='center')

        self._set_table_borders(doc_table)
        self.doc.add_paragraph()

    def _get_alignment(self, alignments: List[str], idx: int) -> str:
        """안전하게 정렬 값 가져오기"""
        if idx < len(alignments):
            return alignments[idx]
        return 'left'

    def _strip_markdown(self, text: str) -> str:
        """마크다운 서식 제거"""
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return text

    def _render_code(self, code: str):
        """코드 블록 렌더링"""
        para = self.doc.add_paragraph()
        self._set_paragraph_shading(para, self.style.colors.bg_info)

        run = para.add_run(code)
        run.font.name = self.style.fonts.code_font
        run.font.size = self.style.fonts.small_size

    def _render_quote(self, content: str):
        """인용구 블록 렌더링 (Back Data 등)"""
        # 인용구 박스로 표현 (info 스타일 배경 + 왼쪽 테두리 효과)
        para = self.doc.add_paragraph()
        self._set_paragraph_shading(para, self.style.callout.info_bg)
        para.paragraph_format.left_indent = Inches(0.2)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)

        # 각 줄을 인라인 서식 처리하여 출력
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if line.strip():
                self._add_formatted_text(para, line.strip())
            if i < len(lines) - 1:
                para.add_run('\n')

    def _render_callout(self, content: Dict, callout_type: str):
        """콜아웃 박스 렌더링"""
        title = content.get('title', '')
        body = content.get('body', '')

        if callout_type == 'info':
            bg_color = self.style.callout.info_bg
            icon = self.style.callout.info_icon
            border_color = self.style.callout.info_border
        else:
            bg_color = self.style.callout.warning_bg
            icon = self.style.callout.warning_icon
            border_color = self.style.callout.warning_border

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = self._content_width

        cell = table.cell(0, 0)
        cell.width = self._content_width

        title_para = cell.paragraphs[0]
        title_text = f"{icon} {title}" if icon else title
        title_run = title_para.add_run(title_text)
        title_run.bold = True
        title_run.font.size = self.style.fonts.body_size

        if body.strip():
            body_para = cell.add_paragraph()
            self._add_formatted_text(body_para, body.strip())

        self._style_cell(cell, bg_color=bg_color)
        self._set_table_borders(table, color=border_color)

        self.doc.add_paragraph()

    def _render_flow(self, flow: ParsedFlow):
        """플로우 다이어그램 렌더링"""
        boxes = [item for item in flow.items if not item.is_arrow]

        if not boxes:
            return

        cols = len(boxes) * 2 - 1
        table = self.doc.add_table(rows=1, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 전체 너비 계산 (동적)
        total_width = self._content_width
        box_width = total_width / cols

        col_idx = 0
        for i, item in enumerate(flow.items):
            cell = table.cell(0, col_idx)
            cell.width = box_width

            if item.is_arrow:
                cell.text = '→'
                self._style_cell(cell, align='center')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(14)
            else:
                cell.text = item.text
                self._style_cell(
                    cell,
                    bg_color=self.style.colors.bg_info,
                    align='center'
                )

            col_idx += 1

        self._set_table_borders(table)

        if flow.caption:
            caption = self.doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption.add_run(flow.caption)
            cap_run.italic = True
            cap_run.font.size = self.style.fonts.caption_size
            cap_run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.text_gray)

        self.doc.add_paragraph()

    def _render_hr(self):
        """수평선 렌더링"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), self.style.colors.border_gray)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_image(self, content: Dict):
        """이미지 렌더링"""
        src = content.get('src', '')
        alt = content.get('alt', '')

        if Path(src).exists():
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(src, width=Inches(4))

            if alt:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt)
                cap_run.italic = True
                cap_run.font.size = self.style.fonts.caption_size

    # =========================================================================
    # 유틸리티
    # =========================================================================

    def _add_page_break(self):
        """페이지 나누기"""
        self.doc.add_page_break()

    def _add_formatted_text(self, para, text: str):
        """인라인 서식이 적용된 텍스트 추가"""
        parts = MarkdownParser.parse_inline(text)

        for part in parts:
            run = para.add_run(part['content'])
            run.font.name = self.style.fonts.default_font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style.fonts.default_font)
            run.font.size = self.style.fonts.body_size

            if part['type'] == 'bold':
                run.bold = True
            elif part['type'] == 'italic':
                run.italic = True
            elif part['type'] == 'code':
                run.font.name = self.style.fonts.code_font
                run.font.size = self.style.fonts.small_size
            elif part['type'] == 'link':
                run.font.color.rgb = self.style.colors.to_rgb(self.style.colors.brand_blue)
                run.underline = True

    # 정렬 매핑
    ALIGN_MAP = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }

    def _style_cell(self, cell, bg_color: str = None, bold: bool = False, align: str = 'left'):
        """셀 스타일 적용"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 배경색
        if bg_color:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)

        # 패딩
        tcMar = OxmlElement('w:tcMar')
        for name, val in [('top', '60'), ('left', '80'), ('bottom', '60'), ('right', '80')]:
            margin = OxmlElement(f'w:{name}')
            margin.set(qn('w:w'), val)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

        # 세로 중앙 정렬 (python-docx API 사용 - 기존 vAlign 제거 후 재설정)
        existing = tcPr.find(qn('w:vAlign'))
        if existing is not None:
            tcPr.remove(existing)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 가로 정렬 + 단락 간격 제거 (세로 중앙 정렬이 정확히 작동하도록)
        word_align = self.ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        for para in cell.paragraphs:
            para.alignment = word_align
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

            for run in para.runs:
                if bold:
                    run.bold = True
                run.font.size = self.style.fonts.body_size
                run.font.name = self.style.fonts.default_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.style.fonts.default_font)

    def _set_table_borders(self, table, color: str = None):
        """테이블 테두리 설정"""
        border_color = color or self.style.table.border_color

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_paragraph_shading(self, para, color: str):
        """단락 배경색 설정"""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        pPr.append(shd)
